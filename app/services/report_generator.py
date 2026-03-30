import datetime
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from jinja2 import Environment, FileSystemLoader
from weasyprint import HTML

from app.branding import (
    COLOR_ACCENT_RGB,
    COLOR_HEADER_MUTED_RGB,
    COLOR_HEADER_RGB,
    DOCX_HEX_HEADER,
    docx_topbar_title,
    report_render_context,
)
from app.models import RankingSession

# ── colour palette (Crossroads / DOCX) ──────────────────────
HEADER_BG    = RGBColor(*COLOR_HEADER_RGB)
ACCENT       = RGBColor(*COLOR_ACCENT_RGB)
HEADER_MUTED = RGBColor(*COLOR_HEADER_MUTED_RGB)
CRIT_PCT     = RGBColor(0x90, 0xBE, 0xF5)   # light accent for % in criteria grid
BADGE_TEXT   = RGBColor(0xDB, 0xEA, 0xFE)   # on dark tags in hero
AMBER        = RGBColor(0xB4, 0x53, 0x09)
RED          = RGBColor(0xB9, 0x1C, 0x1C)
BLUE         = RGBColor(0x1E, 0x40, 0xAF)
GREY         = RGBColor(0x6B, 0x72, 0x80)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
BLACK        = RGBColor(0x1A, 0x1A, 0x1A)
TEXT_NAVY = HEADER_BG
DOCX_HDR  = DOCX_HEX_HEADER


def _match(total: float):
    if total >= 75:  return "TOP MATCH",    ACCENT
    if total >= 55:  return "STRONG MATCH", BLUE
    if total >= 35:  return "CONDITIONAL",  AMBER
    return "POOR FIT", RED


def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a hex colour (no #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, size_pt=9,
               color: RGBColor = None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    para.clear()
    run = para.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _score_color(score: float) -> RGBColor:
    if score >= 8:  return ACCENT
    if score >= 5:  return AMBER
    return RED


class ReportGenerator:
    def __init__(self, templates_dir: str = "app/templates"):
        self.jinja_env = Environment(
            loader=FileSystemLoader(templates_dir),
            autoescape=True,
        )

    def _build_report_context(self, session: RankingSession, top_n: int = 15) -> dict:
        top_candidates = sorted(
            [c for c in session.candidates if c.weighted_total is not None],
            key=lambda c: c.weighted_total,
            reverse=True,
        )[:top_n]

        ctx = report_render_context()
        ctx.update({
            "session": session,
            "job_role": session.job_role,
            "candidates": top_candidates,
            "total_candidates": session.total_candidates,
            "generated_at": datetime.datetime.utcnow(),
            "criteria": session.job_role.criteria,
        })
        return ctx

    # ── PDF ──────────────────────────────────────────────────
    def render_html_report(self, session: RankingSession, top_n: int = 15) -> str:
        template = self.jinja_env.get_template("reports/report_template.html")
        context  = self._build_report_context(session, top_n)
        return template.render(**context)

    def render_pdf_report(self, session: RankingSession, top_n: int = 15) -> bytes:
        return HTML(string=self.render_html_report(session, top_n)).write_pdf()

    # ── DOCX ─────────────────────────────────────────────────
    def render_docx_report(self, session: RankingSession, top_n: int = 15) -> bytes:
        ctx  = self._build_report_context(session, top_n)
        doc  = Document()

        # narrow margins
        for sec in doc.sections:
            sec.top_margin    = Inches(0.65)
            sec.bottom_margin = Inches(0.65)
            sec.left_margin   = Inches(0.75)
            sec.right_margin  = Inches(0.75)

        criteria       = ctx["criteria"]
        crit_names     = [c["name"] for c in criteria]
        all_candidates = sorted(
            [c for c in session.candidates if c.weighted_total is not None],
            key=lambda c: c.weighted_total,
            reverse=True,
        )

        # ════════════════════════════════════════════
        # SECTION 01 — COVER / CRITERIA
        # ════════════════════════════════════════════
        self._add_topbar(doc, ctx)
        self._add_hero(doc, ctx)
        self._section_label(doc, "01", "SCORING CRITERIA", "How candidates were evaluated")
        self._add_criteria_table(doc, criteria)
        doc.add_page_break()

        # ════════════════════════════════════════════
        # SECTION 02 — SCORING MATRIX
        # ════════════════════════════════════════════
        self._section_label(doc, "02", "CANDIDATE SCORING MATRIX",
                            "All candidates · sorted by total score, high to low")
        self._add_matrix_table(doc, all_candidates, criteria)
        doc.add_page_break()

        # ════════════════════════════════════════════
        # SECTION 03 — TOP CANDIDATE PROFILES
        # ════════════════════════════════════════════
        self._section_label(doc, "03", "TOP CANDIDATE PROFILES",
                            f"Top {len(ctx['candidates'])} candidates recommended for recruiter review")
        doc.add_paragraph()

        for candidate in ctx["candidates"]:
            self._add_profile_card(doc, candidate, criteria)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── helpers ──────────────────────────────────────────────

    def _add_topbar(self, doc: Document, ctx: dict):
        p = doc.add_paragraph()
        r = p.add_run(
            f"{docx_topbar_title()}"
            f"                                              "
            f"CONFIDENTIAL  ·  {ctx['generated_at'].strftime('%Y')}"
        )
        r.font.size  = Pt(7.5)
        r.font.color.rgb = HEADER_MUTED
        r.font.bold  = True
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  DOCX_HDR)
        pPr.append(shd)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    def _add_hero(self, doc: Document, ctx: dict):
        # big title
        p = doc.add_paragraph()
        r = p.add_run("CANDIDATE RANKING REPORT")
        r.bold = True
        r.font.size = Pt(28)
        r.font.color.rgb = WHITE
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  DOCX_HDR)
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)

        # subtitle
        p2 = doc.add_paragraph()
        r2 = p2.add_run(
            f"{ctx['job_role'].title}  ·  "
            f"{ctx['generated_at'].strftime('%B %d, %Y')}"
        )
        r2.font.size = Pt(11)
        r2.font.color.rgb = HEADER_MUTED
        pPr2 = p2._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"),   "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"),  DOCX_HDR)
        pPr2.append(shd2)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)

        # badge line
        p3 = doc.add_paragraph()
        for badge_text in [
            "AI-POWERED SCREENING",
            f"{ctx['total_candidates']} CANDIDATES EVALUATED",
            f"TOP {len(ctx['candidates'])} SHOWN",
        ]:
            rb = p3.add_run(f"  {badge_text}  ")
            rb.bold = True
            rb.font.size = Pt(7.5)
            rb.font.color.rgb = BADGE_TEXT
            rb.font.highlight_color = None
            p3.add_run("   ")
        pPr3 = p3._p.get_or_add_pPr()
        shd3 = OxmlElement("w:shd")
        shd3.set(qn("w:val"),   "clear")
        shd3.set(qn("w:color"), "auto")
        shd3.set(qn("w:fill"),  DOCX_HDR)
        pPr3.append(shd3)
        p3.paragraph_format.space_after = Pt(14)

    def _section_label(self, doc: Document, num: str, title: str, sub: str):
        p = doc.add_paragraph()
        rn = p.add_run(f"{num}  ")
        rn.bold = True
        rn.font.size = Pt(18)
        rn.font.color.rgb = RGBColor(0xD1, 0xD5, 0xDB)
        rt = p.add_run(title)
        rt.bold = True
        rt.font.size = Pt(12)
        rt.font.color.rgb = TEXT_NAVY
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(6)
        # border-bottom
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), DOCX_HDR)
        pBdr.append(bottom)
        pPr.append(pBdr)

        p2 = doc.add_paragraph(sub)
        p2.runs[0].font.size = Pt(7.5)
        p2.runs[0].font.color.rgb = GREY
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(10)

    def _add_criteria_table(self, doc: Document, criteria: list):
        letters = list("ABCDEFGH")
        n       = len(criteria)
        table   = doc.add_table(rows=2, cols=n)
        # row 0 — pct + letter header
        for i, crit in enumerate(criteria):
            cell = table.cell(0, i)
            _set_cell_bg(cell, DOCX_HDR)
            cell.paragraphs[0].clear()
            r1 = cell.paragraphs[0].add_run(f"{letters[i]}")
            r1.font.size = Pt(7); r1.font.color.rgb = HEADER_MUTED
            p2 = cell.add_paragraph(f"{crit['weight']}%")
            p2.runs[0].bold = True
            p2.runs[0].font.size = Pt(16)
            p2.runs[0].font.color.rgb = CRIT_PCT
            p3 = cell.add_paragraph(crit["name"])
            p3.runs[0].bold = True
            p3.runs[0].font.size = Pt(7.5)
            p3.runs[0].font.color.rgb = WHITE

        # row 1 — score anchors
        for i in range(n):
            cell = table.cell(1, i)
            cell.paragraphs[0].clear()
            for score_label, desc in [
                ("10", "Exceeds"),
                ("7–9", "Meets fully"),
                ("4–6", "Partial"),
                ("1–3", "Limited"),
                ("0", "No evidence"),
            ]:
                pa = cell.add_paragraph()
                rs = pa.add_run(f"{score_label}  ")
                rs.bold = True; rs.font.size = Pt(7); rs.font.color.rgb = ACCENT
                rd = pa.add_run(desc)
                rd.font.size = Pt(7); rd.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                pa.paragraph_format.space_before = Pt(0)
                pa.paragraph_format.space_after  = Pt(0)

        # clean up empty first para in anchor cells
        for i in range(n):
            cell = table.cell(1, i)
            if not cell.paragraphs[0].text:
                p = cell.paragraphs[0]._element
                p.getparent().remove(p)

        table.style = "Table Grid"
        doc.add_paragraph()

    def _add_matrix_table(self, doc: Document, candidates: list, criteria: list):
        if not candidates:
            doc.add_paragraph("No scored candidates.")
            return

        letters = list("ABCDEFGH")
        n_crit  = len(criteria)
        cols    = 1 + n_crit + 2   # candidate | criteria... | total | match
        table   = doc.add_table(rows=1, cols=cols)
        table.style = "Table Grid"

        # header row
        hdr = table.rows[0].cells
        _set_cell_bg(hdr[0], DOCX_HDR)
        _cell_text(hdr[0], "CANDIDATE", bold=True, size_pt=7.5, color=WHITE)

        for i, crit in enumerate(criteria):
            _set_cell_bg(hdr[i + 1], DOCX_HDR)
            _cell_text(hdr[i + 1], f"{letters[i]} – {crit['name']}\n{crit['weight']}%",
                       bold=True, size_pt=7, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

        _set_cell_bg(hdr[-2], DOCX_HDR)
        _cell_text(hdr[-2], "TOTAL\n/ 100", bold=True, size_pt=7.5, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_bg(hdr[-1], DOCX_HDR)
        _cell_text(hdr[-1], "MATCH", bold=True, size_pt=7.5, color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

        # data rows
        for candidate in candidates:
            row   = table.add_row().cells
            total = candidate.weighted_total
            match_label, match_color = _match(total)

            # candidate cell
            cand_para = row[0].paragraphs[0]
            cand_para.clear()
            rn = cand_para.add_run(f"#{candidate.rank}  {candidate.candidate_name or 'Unknown'}")
            rn.bold = True; rn.font.size = Pt(8)
            sub_para = row[0].add_paragraph(candidate.filename or "")
            sub_para.runs[0].font.size = Pt(6.5)
            sub_para.runs[0].font.color.rgb = GREY

            # criterion scores
            crit_map = {cs.criterion_name: cs for cs in (candidate.criterion_scores or [])}
            for i, crit in enumerate(criteria):
                cs = crit_map.get(crit["name"])
                s  = cs.score if cs else 0
                _cell_text(row[i + 1],
                           f"{s:.0f}" if cs else "—",
                           bold=True, size_pt=9,
                           color=_score_color(s) if cs else GREY,
                           align=WD_ALIGN_PARAGRAPH.CENTER)

            # total
            if total >= 75:   tot_color = ACCENT
            elif total >= 55: tot_color = AMBER
            else:             tot_color = RED
            _cell_text(row[-2], f"{total:.1f}", bold=True, size_pt=10,
                       color=tot_color, align=WD_ALIGN_PARAGRAPH.CENTER)

            # match
            _cell_text(row[-1], match_label, bold=True, size_pt=7,
                       color=match_color, align=WD_ALIGN_PARAGRAPH.CENTER)

        # column widths
        cand_w  = Inches(1.8)
        total_w = Inches(0.55)
        match_w = Inches(0.8)
        page_w  = Inches(7.0)
        crit_w  = (page_w - cand_w - total_w - match_w) / n_crit
        for row in table.rows:
            row.cells[0].width = cand_w
            for i in range(n_crit):
                row.cells[i + 1].width = crit_w
            row.cells[-2].width = total_w
            row.cells[-1].width = match_w

        doc.add_paragraph()

    def _add_profile_card(self, doc: Document, candidate, criteria: list):
        total = candidate.weighted_total
        match_label, match_color = _match(total)
        crit_map = {cs.criterion_name: cs for cs in (candidate.criterion_scores or [])}
        raw   = candidate.raw_scores or {}
        flags = raw.get("red_flags", []) if isinstance(raw, dict) else []

        # ── rank + name header row (2-col table) ──
        hdr_table = doc.add_table(rows=1, cols=2)
        left  = hdr_table.cell(0, 0)
        right = hdr_table.cell(0, 1)
        hdr_table.cell(0, 0).width = Inches(1.1)
        hdr_table.cell(0, 1).width = Inches(5.9)

        # left: rank + score
        left.paragraphs[0].clear()
        rk = left.paragraphs[0].add_run(f"Rank {candidate.rank}")
        rk.font.size = Pt(7); rk.font.color.rgb = GREY
        sc_p = left.add_paragraph(f"{total:.1f}")
        sc_p.runs[0].bold = True; sc_p.runs[0].font.size = Pt(22)
        sc_p.runs[0].font.color.rgb = TEXT_NAVY
        sc_p.paragraph_format.space_before = Pt(0)
        sc_p.paragraph_format.space_after  = Pt(0)
        den_p = left.add_paragraph("/ 100 pts")
        den_p.runs[0].font.size = Pt(7); den_p.runs[0].font.color.rgb = GREY
        den_p.paragraph_format.space_before = Pt(0)
        den_p.paragraph_format.space_after  = Pt(4)
        mtch_p = left.add_paragraph(match_label)
        mtch_p.runs[0].bold = True; mtch_p.runs[0].font.size = Pt(7)
        mtch_p.runs[0].font.color.rgb = match_color
        mtch_p.paragraph_format.space_before = Pt(0)

        # right: name + filename + skill tags
        right.paragraphs[0].clear()
        nm = right.paragraphs[0].add_run(
            (candidate.candidate_name or "Unknown Candidate").upper()
        )
        nm.bold = True; nm.font.size = Pt(14); nm.font.color.rgb = TEXT_NAVY
        fn_p = right.add_paragraph(candidate.filename or "")
        fn_p.runs[0].font.size = Pt(7.5); fn_p.runs[0].font.color.rgb = GREY
        fn_p.paragraph_format.space_before = Pt(2); fn_p.paragraph_format.space_after = Pt(4)

        # skill tags = criteria scored ≥ 7
        tags = [c["name"] for c in criteria
                if crit_map.get(c["name"]) and crit_map[c["name"]].score >= 7]
        if tags:
            tag_p = right.add_paragraph()
            for tag in tags:
                rt = tag_p.add_run(f"  {tag}  ")
                rt.bold = True; rt.font.size = Pt(7)
                rt.font.color.rgb = BADGE_TEXT
            tag_p.paragraph_format.space_after = Pt(6)
            pPr = tag_p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  DOCX_HDR)
            pPr.append(shd)

        doc.add_paragraph()

        # ── summary ──
        if candidate.ai_summary:
            sp = doc.add_paragraph(candidate.ai_summary)
            sp.runs[0].font.size = Pt(8.5)
            sp.paragraph_format.space_after = Pt(6)

        # ── criterion scores table ──
        n_crit  = len(criteria)
        sc_tbl  = doc.add_table(rows=1, cols=4)
        sc_tbl.style = "Table Grid"
        for cell, label in zip(sc_tbl.rows[0].cells,
                               ["Criterion", "Weight", "Score", "Justification"]):
            _cell_text(cell, label, bold=True, size_pt=7.5,
                       color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            _set_cell_bg(cell, DOCX_HDR)

        for crit in criteria:
            cs  = crit_map.get(crit["name"])
            s   = cs.score if cs else 0
            row = sc_tbl.add_row().cells
            _cell_text(row[0], crit["name"], size_pt=8)
            _cell_text(row[1], f"{crit['weight']}%", size_pt=8,
                       color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row[2], f"{s:.0f}/10" if cs else "—",
                       bold=True, size_pt=8,
                       color=_score_color(s) if cs else GREY,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell_text(row[3], cs.justification if cs else "", size_pt=7.5)

        # column widths
        for row in sc_tbl.rows:
            row.cells[0].width = Inches(1.4)
            row.cells[1].width = Inches(0.5)
            row.cells[2].width = Inches(0.55)
            row.cells[3].width = Inches(4.55)

        # ── red flags / concerns ──
        if flags:
            doc.add_paragraph()
            fp = doc.add_paragraph()
            fr = fp.add_run("Concerns:  ")
            fr.bold = True; fr.font.size = Pt(8); fr.font.color.rgb = AMBER
            fp.add_run("  ·  ".join(flags)).font.size = Pt(8)
            pPr = fp._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "FFFBEB")
            pPr.append(shd)

        # spacer + thin rule
        sp2 = doc.add_paragraph()
        sp2.paragraph_format.space_before = Pt(8)
        sp2.paragraph_format.space_after  = Pt(0)
        pPr = sp2._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E5E7EB")
        pBdr.append(bottom)
        pPr.append(pBdr)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
