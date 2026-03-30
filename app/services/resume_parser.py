from dataclasses import dataclass
from pathlib import Path

import pdfplumber
from docx import Document

MIN_TEXT_LENGTH = 50
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


@dataclass
class ParsedResume:
    filename: str
    text: str
    page_count: int
    error: str | None = None


class ResumeParser:

    def discover_files(self, folder_path: Path) -> list[Path]:
        """Return all PDF and DOCX files in the folder."""
        if not folder_path.is_dir():
            raise FileNotFoundError(f"Resume folder not found: {folder_path}")

        files = []
        for ext in SUPPORTED_EXTENSIONS:
            files.extend(folder_path.glob(f"*{ext}"))
        return sorted(files, key=lambda p: p.name.lower())

    def parse_file(self, file_path: Path) -> ParsedResume:
        """Extract text from a single resume file."""
        ext = file_path.suffix.lower()
        try:
            if ext == ".pdf":
                return self._parse_pdf(file_path)
            elif ext == ".docx":
                return self._parse_docx(file_path)
            elif ext == ".txt":
                return self._parse_txt(file_path)
            else:
                return ParsedResume(
                    filename=file_path.name, text="", page_count=0,
                    error=f"Unsupported file type: {ext}",
                )
        except Exception as e:
            return ParsedResume(
                filename=file_path.name, text="", page_count=0,
                error=f"Parse error: {e}",
            )

    def _parse_pdf(self, file_path: Path) -> ParsedResume:
        with pdfplumber.open(file_path) as pdf:
            pages = []
            for page in pdf.pages:
                text = page.extract_text(layout=True)
                if text:
                    pages.append(text)
            full_text = "\n\n".join(pages)

            if len(full_text.strip()) < MIN_TEXT_LENGTH:
                return ParsedResume(
                    filename=file_path.name,
                    text=full_text.strip(),
                    page_count=len(pdf.pages),
                    error="Extracted text is too short. This may be a scanned/image-only PDF.",
                )

            return ParsedResume(
                filename=file_path.name,
                text=full_text.strip(),
                page_count=len(pdf.pages),
            )

    def _parse_docx(self, file_path: Path) -> ParsedResume:
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables (common in formatted resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        full_text = "\n".join(paragraphs)

        if len(full_text.strip()) < MIN_TEXT_LENGTH:
            return ParsedResume(
                filename=file_path.name,
                text=full_text.strip(),
                page_count=1,
                error="Extracted text is too short. Document may be empty or image-based.",
            )

        return ParsedResume(
            filename=file_path.name,
            text=full_text.strip(),
            page_count=1,
        )

    def _parse_txt(self, file_path: Path) -> ParsedResume:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        if len(text.strip()) < MIN_TEXT_LENGTH:
            return ParsedResume(
                filename=file_path.name,
                text=text.strip(),
                page_count=1,
                error="Extracted text is too short.",
            )
        return ParsedResume(filename=file_path.name, text=text.strip(), page_count=1)

    def parse_folder(self, folder_path: Path) -> list[ParsedResume]:
        """Parse all resumes in a folder."""
        files = self.discover_files(folder_path)
        return [self.parse_file(f) for f in files]
